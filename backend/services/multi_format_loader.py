"""多格式文档加载器

支持 DOCX / Excel / TXT / Markdown 格式，
输出与 PDF 提取结果兼容的数据结构。

参考 kotaemon 的 docx_loader / excel_loader / html_loader。
"""

import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# 支持的文件扩展名
SUPPORTED_EXTENSIONS = {".docx", ".xlsx", ".xls", ".txt", ".md", ".markdown", ".csv"}


def is_supported_format(filename: str) -> bool:
    """检查文件格式是否受支持"""
    return Path(filename).suffix.lower() in SUPPORTED_EXTENSIONS


def extract_from_file(filepath: str, filename: str) -> dict:
    """从非 PDF 文件中提取文本

    Args:
        filepath: 文件绝对路径
        filename: 原始文件名

    Returns:
        与 PDF 提取结果兼容的字典:
        {
            "full_text": str,
            "total_pages": int,
            "pages": [{"page": int, "text": str}, ...],
            "source_type": str,
        }
    """
    ext = Path(filename).suffix.lower()

    if ext == ".docx":
        return _extract_docx(filepath)
    elif ext in (".xlsx", ".xls"):
        return _extract_excel(filepath)
    elif ext == ".csv":
        return _extract_csv(filepath)
    elif ext in (".txt", ".md", ".markdown"):
        return _extract_text(filepath, ext)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


def _extract_docx(filepath: str) -> dict:
    """从 DOCX 文件提取文本

    使用 python-docx 解析段落和表格。
    表格转为 Markdown 格式。
    """
    try:
        from docx import Document
    except ImportError:
        raise ValueError("需要安装 python-docx: pip install python-docx")

    doc = Document(filepath)
    parts = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # 段落
            for para in doc.paragraphs:
                if para._element == element:
                    text = para.text.strip()
                    if text:
                        # 保留标题样式
                        if para.style and para.style.name.startswith("Heading"):
                            level = para.style.name.replace("Heading ", "").strip()
                            try:
                                level = int(level)
                            except ValueError:
                                level = 1
                            parts.append(f"{'#' * level} {text}")
                        else:
                            parts.append(text)
                    break

        elif tag == "tbl":
            # 表格
            for table in doc.tables:
                if table._element == element:
                    md_table = _table_to_markdown(table)
                    if md_table:
                        parts.append(md_table)
                    break

    full_text = "\n\n".join(parts)

    # 按段落分页（每 ~3000 字符为一"页"）
    pages = _split_to_pages(full_text, chars_per_page=3000)

    logger.info(f"[DocxLoader] 提取 {len(parts)} 个元素, {len(full_text)} 字符")
    return {
        "full_text": full_text,
        "total_pages": len(pages),
        "pages": pages,
        "source_type": "docx",
    }


def _table_to_markdown(table) -> str:
    """将 python-docx 表格转为 Markdown 格式"""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")

    if len(rows) < 1:
        return ""

    # 在第一行后插入分隔行
    header = rows[0]
    num_cols = header.count("|") - 1
    separator = "| " + " | ".join(["---"] * num_cols) + " |"

    return "\n".join([rows[0], separator] + rows[1:])


def _extract_excel(filepath: str) -> dict:
    """从 Excel 文件提取文本

    使用 openpyxl 解析多 sheet。
    每个 sheet 转为 Markdown 表格。
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ValueError("需要安装 openpyxl: pip install openpyxl")

    wb = load_workbook(filepath, read_only=True, data_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"## {sheet_name}")

        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                rows.append("| " + " | ".join(cells) + " |")

        if rows:
            num_cols = rows[0].count("|") - 1
            separator = "| " + " | ".join(["---"] * max(num_cols, 1)) + " |"
            parts.append("\n".join([rows[0], separator] + rows[1:]))

    wb.close()
    full_text = "\n\n".join(parts)
    pages = _split_to_pages(full_text, chars_per_page=3000)

    logger.info(f"[ExcelLoader] {len(wb.sheetnames)} sheets, {len(full_text)} 字符")
    return {
        "full_text": full_text,
        "total_pages": len(pages),
        "pages": pages,
        "source_type": "excel",
    }


def _extract_csv(filepath: str) -> dict:
    """从 CSV 文件提取文本"""
    import csv

    parts = []
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        rows = []
        for row in reader:
            cells = [c.strip() for c in row]
            rows.append("| " + " | ".join(cells) + " |")

    if rows:
        num_cols = rows[0].count("|") - 1
        separator = "| " + " | ".join(["---"] * max(num_cols, 1)) + " |"
        parts.append("\n".join([rows[0], separator] + rows[1:]))

    full_text = "\n\n".join(parts)
    pages = _split_to_pages(full_text, chars_per_page=3000)

    logger.info(f"[CSVLoader] {len(rows)} 行, {len(full_text)} 字符")
    return {
        "full_text": full_text,
        "total_pages": len(pages),
        "pages": pages,
        "source_type": "csv",
    }


def _extract_text(filepath: str, ext: str) -> dict:
    """从纯文本或 Markdown 文件提取文本"""
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        full_text = f.read()

    source_type = "markdown" if ext in (".md", ".markdown") else "text"
    pages = _split_to_pages(full_text, chars_per_page=3000)

    logger.info(f"[TextLoader] {len(full_text)} 字符, type={source_type}")
    return {
        "full_text": full_text,
        "total_pages": len(pages),
        "pages": pages,
        "source_type": source_type,
    }


def _split_to_pages(text: str, chars_per_page: int = 3000) -> list[dict]:
    """将长文本按字符数分割为"页"

    尽量在段落边界处分割。

    Args:
        text: 输入文本
        chars_per_page: 每页近似字符数

    Returns:
        [{"page": int, "text": str}, ...]
    """
    if not text:
        return [{"page": 1, "text": ""}]

    paragraphs = text.split("\n\n")
    pages = []
    current_page_text = ""
    page_num = 1

    for para in paragraphs:
        if len(current_page_text) + len(para) > chars_per_page and current_page_text:
            pages.append({"page": page_num, "text": current_page_text.strip()})
            page_num += 1
            current_page_text = para + "\n\n"
        else:
            current_page_text += para + "\n\n"

    if current_page_text.strip():
        pages.append({"page": page_num, "text": current_page_text.strip()})

    return pages if pages else [{"page": 1, "text": text}]
